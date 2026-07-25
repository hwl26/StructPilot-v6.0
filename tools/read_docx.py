"""
读取 .docx 文件的工具脚本

.docx 是压缩的 XML 格式，可以通过 python-docx 库读取
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("❌ 需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def read_docx(file_path: str) -> None:
    """读取 docx 文件并输出内容"""
    doc_path = Path(file_path)

    if not doc_path.exists():
        print(f"❌ 文件不存在：{file_path}")
        return

    try:
        doc = Document(doc_path)

        print(f"📄 文件：{doc_path.name}")
        print(f"📊 段落数：{len(doc.paragraphs)}")
        print(f"🖼️ 图片数：{len([r for p in doc.paragraphs for r in p.runs if r._element.xpath('.//pic:pic')])}")
        print("=" * 60)
        print()

        # 读取所有段落
        for i, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                print(f"[{i}] {text}")
                print()

        # 读取表格
        if doc.tables:
            print("=" * 60)
            print(f"📊 表格数：{len(doc.tables)}")
            print()

            for table_idx, table in enumerate(doc.tables, 1):
                print(f"表格 {table_idx}:")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    print(" | ".join(cells))
                print()

    except Exception as e:
        print(f"❌ 读取失败：{e}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python read_docx.py <文件路径>")
        print("示例：python read_docx.py document.docx")
        sys.exit(1)

    file_path = sys.argv[1]
    read_docx(file_path)
