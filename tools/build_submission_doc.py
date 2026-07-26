from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "比赛提交文档" / "StructPilot_v6_决赛作品说明文档_提交版.docx"
RED = "A52A2A"
RED_DARK = "7F1D1D"
RED_LIGHT = "F9E8E6"
INK = "242124"
MUTED = "6B6462"
LINE = "D9D2CF"
PALE = "FBF8F7"
GREEN = "2F6B54"
AMBER = "8A5A00"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in kwargs[edge]:
                element.set(qn("w:" + key), str(kwargs[edge][key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("StructPilot v6.0  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_run(p, text, bold=False, color=None, size=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return r


def add_body(doc, text, *, first_line=True, space_after=6, color=INK):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_run(p, text, color=color)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55 + 0.45 * level)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        add_run(p, item)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, bold=True, color=RED_DARK if level == 1 else RED, size=16 if level == 1 else (12.5 if level == 2 else 11.2))
    return p


def add_label(doc, label, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(4)
    add_run(p, label + "  ", bold=True, color=RED)
    add_run(p, text)
    return p


def add_callout(doc, title, text, fill=RED_LIGHT, accent=RED):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.8)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": accent}, top={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=accent, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, text, color=INK, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, header_fill=RED_DARK, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        if widths:
            cell.width = Cm(widths[i])
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, 90, 95, 90, 95)
        set_cell_border(cell, top={"val": "single", "sz": "6", "color": header_fill}, bottom={"val": "single", "sz": "6", "color": header_fill}, left={"val": "single", "sz": "4", "color": "FFFFFF"}, right={"val": "single", "sz": "4", "color": "FFFFFF"})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, bold=True, color="FFFFFF", size=font_size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if widths:
                cell.width = Cm(widths[i])
            set_cell_shading(cell, "FFFFFF" if ridx % 2 == 0 else PALE)
            set_cell_margins(cell, 80, 95, 80, 95)
            set_cell_border(cell, top={"val": "single", "sz": "3", "color": LINE}, bottom={"val": "single", "sz": "3", "color": LINE}, left={"val": "single", "sz": "3", "color": LINE}, right={"val": "single", "sz": "3", "color": LINE})
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            add_run(p, str(value), color=INK, size=font_size)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path: Path, caption: str, width_cm=15.8):
    if not path.exists():
        add_callout(doc, "图示缺失", f"预期图片未找到：{path.name}", fill="FFF4E5", accent=AMBER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(7)
    add_run(cp, caption, color=MUTED, size=8.5, italic=True)


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "资料依据：", bold=True, color=MUTED, size=8.5)
    add_run(p, text, color=MUTED, size=8.5)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for name, size, color in (("Heading 1", 16, RED_DARK), ("Heading 2", 12.5, RED), ("Heading 3", 11.2, RED)):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.2)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(header, "STRUCTPILOT  /  冷冻电镜单颗粒分析智能导航系统", bold=True, color=RED, size=8.2)
    header.paragraph_format.space_after = Pt(0)
    p = sec.footer.paragraphs[0]
    add_page_number(p)


def add_cover(doc):
    cover = ROOT / "01_cover.png"
    if cover.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        p.add_run().add_picture(str(cover), width=Cm(17.0))
    else:
        add_heading(doc, "StructPilot v6.0", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "决赛作品说明文档", bold=True, color=RED_DARK, size=20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_run(p, "供评审参阅材料 · 版本 6.0.0", color=MUTED, size=10.5)
    add_callout(doc, "一句话定位", "StructPilot 是面向 cryo-EM 单颗粒分析的流程导航与知识协作系统：把分散的操作步骤、参数含义、质量控制和故障排查组织成可追踪的对话式工作流。", fill=RED_LIGHT, accent=RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    add_run(p, "提交版本：StructPilot_v6 · 2026 年 7 月", color=MUTED, size=9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "团队编号：b-2", color=MUTED, size=9)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "目录与阅读说明", 1)
    add_body(doc, "本文档以“创意—功能—实现—演示—边界”为主线，帮助评审快速理解 StructPilot 解决的问题、实际可用的能力以及当前版本的工程边界。表格中的“已实现”以当前代码和运行界面为准；“需配置”表示能力已预留但依赖外部模型或服务；“边界”用于提示不能替代专业软件或实验判断的部分。")
    items = [
        "01 作品概览：为什么需要一个 cryo-EM 流程导航系统",
        "02 创意说明：把隐性经验变成可追踪的工作流",
        "03 目标用户与应用场景",
        "04 目标功能与 12 个检查点",
        "05 三种交互模式：同一底座，不同认知负担",
        "06 技术架构与知识治理",
        "07 演示示例：从首次进入到参数工作流导出",
        "08 创新点、可靠性与隐私设计",
        "09 部署要求、当前边界与后续路线",
        "附录：功能核验表与建议演示脚本",
    ]
    add_numbered(doc, items)
    add_callout(doc, "评审提示", "StructPilot 的价值不在于替代 cryoSPARC 或 RELION 的计算引擎，而在于降低流程认知成本、减少漏检与误设、沉淀团队经验，并让每一步的“为什么”和“下一步”变得可见。", fill="F3F5F4", accent=GREEN)
    doc.add_page_break()


def add_overview(doc):
    add_heading(doc, "01 作品概览", 1)
    add_heading(doc, "1.1 项目摘要", 2)
    add_body(doc, "冷冻电镜单颗粒分析通常由多个软件、多个参数和多轮质量判断组成。对于初学者，真正的难点不是“点击哪个按钮”，而是理解当前阶段的输入输出、知道哪些参数必须与采集条件一致、能够识别质量问题并决定是否回退。StructPilot 将这些跨软件、跨阶段的知识组织为一套对话式导航系统，服务于 cryoSPARC 与 RELION 工作流。")
    add_body(doc, "当前版本以 Streamlit 提供 Web 界面，以 LangGraph 组织状态和路由，以本地 JSON/Markdown 知识库、规则库和 SQLite 会话存储作为可追踪底座。用户可以从当前检查点开始，查看 SOP、参数解释、QC 提醒、常见故障路径和可选的教学内容；在配置模型服务后，还可以启用文本改写、图像理解和语音转写等增强能力。")
    add_table(doc, ["项目项", "当前版本事实"], [
        ("作品名称", "StructPilot v6.0.0"),
        ("产品定位", "Cryo-EM 单颗粒分析流程导航与知识协作系统"),
        ("界面形态", "本地 Streamlit Web UI"),
        ("支持平台", "cryoSPARC、RELION（以流程指导和知识导航为主）"),
        ("流程结构", "12 个标准检查点，支持进行中、完成、跳过、失败等状态"),
        ("持久化", "SQLite 保存会话、消息和检查点记录"),
        ("增强能力", "可选 LLM、Embedding、Vision、Speech 配置；未配置时保留规则与知识库模式"),
        ("输出", "报告/Markdown 下载；cryoSPARC Workflow JSON 导出"),
    ], widths=[4.0, 12.5])
    add_heading(doc, "1.2 核心价值", 2)
    add_bullets(doc, [
        "把从数据导入到模型验证的长流程拆成可见、可回溯的检查点。",
        "把参数推荐、质量控制和故障排查放在同一上下文中，减少在教程、聊天记录和软件界面之间来回切换。",
        "通过入门、教学、专家三种模式，让同一套底层知识适配不同经验水平。",
        "以本地规则和知识库为最低可用能力，外部模型作为增强而非单点依赖。",
    ])
    add_source_note(doc, "version.py、main.py、graph/、knowledge_base/、runtime/ 与当前运行界面核对。")
    doc.add_page_break()


def add_creative(doc):
    add_heading(doc, "02 创意说明", 1)
    add_heading(doc, "2.1 从“软件操作”转向“实验决策陪跑”", 2)
    add_body(doc, "传统软件界面擅长执行计算，却不一定解释决策。用户面对一个参数时，往往同时需要回答三个问题：它在当前阶段解决什么问题？应该结合哪些采集条件或上游结果设置？运行后如何判断结果是否可信？StructPilot 的创意是把这三个问题固定为每个检查点的最小决策单元：做什么、为什么、怎么验。")
    add_callout(doc, "设计命题", "让流程状态成为对话的上下文，让知识条目成为回答的证据，让质量控制成为下一步行动的入口。", fill=RED_LIGHT, accent=RED)
    add_heading(doc, "2.2 三个设计原则", 2)
    add_table(doc, ["原则", "在软件中的体现", "评审可观察结果"], [
        ("过程可见", "12 个检查点 + 当前阶段 + 完成/失败/跳过状态", "用户能知道自己在哪一步、下一步是什么"),
        ("证据优先", "结构化知识库、规则、QC 标准、来源和审查目录", "回答可回溯，模型输出不是唯一依据"),
        ("渐进增强", "无 API 时走本地规则；配置后启用 LLM/RAG/Vision/Speech", "核心流程不因外部服务不可用而整体失效"),
    ], widths=[3.0, 7.4, 6.1])
    add_heading(doc, "2.3 与普通聊天机器人的区别", 2)
    add_body(doc, "StructPilot 不是把问题丢给一个通用聊天窗口，而是先识别当前软件、检查点、问题类型和流程状态，再调用对应的导航、专家、SOP、记忆或知识检索能力。对于“开始、完成、跳过、报错”等操作性输入，系统会更新检查点状态并留下记录；对于“参数怎么设、这一步是什么、为什么失败”等问题，则返回分层解释和下一步建议。")
    add_source_note(doc, "依据 graph/app.py、agents/、knowledge_base/flows/ 与 knowledge_base/rules/ 的当前实现整理。")
    doc.add_page_break()


def add_users(doc):
    add_heading(doc, "03 目标用户与应用场景", 1)
    add_table(doc, ["用户", "典型困难", "StructPilot 提供的支持"], [
        ("冷冻电镜初学者", "不熟悉流程顺序、术语和基本 QC", "入门模式、分步 SOP、参数解释、教学卡片与测验"),
        ("实验室分析成员", "需要在多个项目间复用经验并保持流程一致", "会话持久化、检查点记录、报告下载、知识沉淀"),
        ("有经验的研究者", "希望快速核对参数边界、比较方案并导出配置", "专家模式、规则检索、参数对照、cryoSPARC Workflow JSON"),
        ("知识库管理员", "SOP 和经验需要审核、分级和持续更新", "导入、审查、验证和 review 目录；角色区分"),
    ], widths=[3.0, 6.0, 7.5])
    add_heading(doc, "3.1 三类应用场景", 2)
    add_numbered(doc, [
        "首次上手：选择 cryoSPARC 或 RELION，确定当前检查点，以入门模式按步骤完成操作并查看 QC 提醒。",
        "问题排查：上传 CTF、2D class、FSC 或 refinement 截图，结合文字描述定位可能原因；未配置 Vision 时仍可使用文字与规则库排查。",
        "团队复用：把一次会话中的参数、判断和结果导出为报告或工作流配置，作为后续实验的参考，而不是直接替代实验员审批。",
    ])
    add_callout(doc, "适用范围", "当前实现适合本地演示、个人研究和小团队流程协作；SQLite 本地存储不是面向大规模多用户生产环境的数据库方案。", fill="FFF4E5", accent=AMBER)
    doc.add_page_break()


def add_functions(doc):
    add_heading(doc, "04 目标功能与 12 个检查点", 1)
    add_heading(doc, "4.1 目标功能清单", 2)
    add_table(doc, ["功能域", "功能描述", "实现状态"], [
        ("流程导航", "选择软件体系、定位当前检查点、记录开始/完成/跳过/失败", "已实现"),
        ("操作指导", "按阶段输出步骤、输入、输出、关键参数和常见陷阱", "已实现"),
        ("质量控制", "提供 QC 检查清单、风险提醒与回退建议", "已实现（知识库/规则驱动）"),
        ("教学支持", "教学卡片、参数含义、常见问题与互动测验", "已实现（部分检查点有专门卡片）"),
        ("专家工具", "参数解释、实验室经验对照、预设管理、配置导出", "已实现"),
        ("多模态输入", "文本、图片上传/粘贴、可选语音转写", "图片/文本已实现；语音需配置服务"),
        ("知识治理", "知识导入、索引、校验、审核目录和状态管理", "已实现"),
        ("会话与报告", "SQLite 保存会话与消息，支持恢复、报告和 Markdown 下载", "已实现"),
        ("工作流导出", "生成 cryoSPARC 可导入的 Workflow JSON", "已实现；不自动提交作业"),
    ], widths=[3.0, 9.8, 3.7])
    add_heading(doc, "4.2 标准分析链路", 2)
    checkpoints = [
        ("01", "数据导入", "导入 movies/micrographs，核对 pixel size、电压、Cs、总剂量"),
        ("02", "运动校正", "校正帧间漂移，检查 drift 曲线和输出路径"),
        ("03", "CTF 估计与曝光筛选", "估计 CTF，剔除冰层、污染和拟合质量差的曝光"),
        ("04", "颗粒挑选", "在 Blob、Topaz、Template 等方法间选择并检查 picks"),
        ("05", "颗粒提取", "确定 box size、Fourier crop 与粒子数量，生成提取集"),
        ("06", "2D 分类与筛选", "通过 class averages 识别清晰类并筛选粒子"),
        ("07", "Ab-initio 初始模型", "从 2D 粒子反推一个或多个初始 3D 模型"),
        ("08", "3D 分类/异质精修", "区分混合构象或不同组分，比较各 class 质量"),
        ("09", "3D 精修", "执行 homogeneous 与 non-uniform refinement，查看 FSC"),
        ("10", "CTF 精修", "全局与局部修正 beam tilt、像散和逐粒子 defocus 等"),
        ("11", "后处理与锐化", "依据 half maps、mask 和 FSC 生成最终密度图"),
        ("12", "模型构建与验证", "在 Coot/Phenix/ChimeraX 等外部工具中建模与验证"),
    ]
    add_table(doc, ["编号", "检查点", "阶段目标"], checkpoints, widths=[1.2, 4.3, 11.0], font_size=8.7)
    add_source_note(doc, "12 个检查点名称和顺序来自 knowledge_base/flows/pipeline_checkpoints.json；各软件具体参数需结合样品、采集条件和当前版本界面复核。")
    doc.add_page_break()


def add_modes(doc):
    add_heading(doc, "05 三种交互模式", 1)
    add_body(doc, "三种模式共享同一套底层流程状态、知识库和编排逻辑，区别主要在信息密度和交互节奏。这样既不牺牲系统一致性，也避免初学者被专家参数淹没。")
    add_table(doc, ["模式", "面向人群", "界面与交互", "适合任务"], [
        ("入门 / 快速模式", "首次使用者、需要按步骤操作的人", "突出当前步骤、操作指令、关键参数和质量检查；减少一次性信息量", "按 SOP 完成一个检查点"),
        ("教学模式", "希望理解原理的学习者", "教学卡片解释“做什么、参数含义、常见问题、判断标准”，配合测验", "理解 2D、CTF、精修等关键概念"),
        ("专家模式", "有经验的研究者、项目负责人", "保留高信息密度工作区，支持参数、规则、经验和导出工具", "核对参数、比较方案、复用预设、导出工作流"),
    ], widths=[3.2, 4.0, 6.2, 3.1])
    add_heading(doc, "5.1 界面证据", 2)
    add_figure(doc, ROOT / "03_beginner.png", "图 1  入门模式视觉示意：把当前步骤、操作指令和 QC 反馈放在同一视野。", width_cm=15.8)
    add_figure(doc, ROOT / "04_teaching.png", "图 2  教学模式视觉示意：通过卡片和测验解释关键概念。", width_cm=15.8)
    add_figure(doc, ROOT / "05_expert.png", "图 3  专家模式视觉示意：支持参数、经验和工作流导出。", width_cm=15.8)
    doc.add_page_break()


def add_architecture(doc):
    add_heading(doc, "06 技术架构与知识治理", 1)
    add_heading(doc, "6.1 分层架构", 2)
    add_table(doc, ["层级", "主要组件", "责任"], [
        ("界面层", "Streamlit Web UI、模式视图、上传与下载组件", "承载对话、检查点、报告和模式切换"),
        ("编排层", "LangGraph、StructPilotApp、路由与状态机", "把输入路由到合适的能力，并维护流程状态"),
        ("智能能力层", "Navigator、Expert、SOP、Memory、QA/绘图相关能力", "提供导航、参数解释、步骤输出、记忆和辅助分析"),
        ("知识层", "流程 JSON、规则、SOP、故障库、教学卡片、审查目录", "提供结构化证据、检索和可更新内容"),
        ("持久化层", "SQLite、缓存与运行目录", "保存会话、消息、检查点记录和部分中间结果"),
    ], widths=[3.0, 6.0, 7.5])
    add_heading(doc, "6.2 单次问答的处理路径", 2)
    add_numbered(doc, [
        "界面接收文本、图片或音频输入，并记录当前软件与检查点。",
        "Navigator 识别意图：概念、参数、故障、进度或通用问答。",
        "规则与知识库先给出结构化结论；必要时检索相关条目。",
        "若配置 LLM/Vision/Speech，则在结构化结论之上进行语言改写、图像理解或转写。",
        "Memory 将消息与状态写入 SQLite，界面返回可执行的下一步和 QC 提醒。",
    ])
    add_heading(doc, "6.3 知识治理闭环", 2)
    add_body(doc, "知识不是一次性写入的静态附件。当前仓库包含导入、索引、校验、review、正式答案、待审核和废弃条目等目录或文件，支持把实验室 SOP、官方文档和经过审查的经验分层管理。系统的基本策略是：结构化规则负责确定性结论，知识条目提供背景与证据，外部模型只负责在可控上下文中增强表达。")
    add_callout(doc, "可靠性原则", "当 Embedding、LLM、Vision 或 Speech 未配置、不可用或调用失败时，系统仍可使用本地规则、流程卡片和知识库完成基础导航；但模型增强相关能力应明确标注为“需配置”。", fill="F3F5F4", accent=GREEN)
    add_source_note(doc, "依据 graph/、agents/、knowledge_base/review/、knowledge_base/rules/ 和 runtime/ 的当前目录结构与实现。")
    doc.add_page_break()


def add_demo(doc):
    add_heading(doc, "07 演示示例", 1)
    add_heading(doc, "7.1 示例 A：首次进入与流程规划", 2)
    add_body(doc, "演示目标：展示系统如何从“选择软件—选择模式—进入检查点”建立上下文，而不是直接给出一段脱离阶段的答案。")
    add_numbered(doc, [
        "打开本地 Web UI，选择 cryoSPARC 或 RELION。",
        "选择入门/快速、教学或专家模式；系统保留同一流程状态。",
        "在当前检查点查看阶段目标、输入、操作步骤、关键参数与 QC 清单。",
        "输入“开始”“完成”“跳过”或“有报错”，检查点状态随之更新，并写入会话记录。",
    ])
    add_figure(doc, ROOT / "比赛提交文档" / "fig_live_home.png", "图 4  真实运行界面：当前位于“步骤 3 · CTF 估计与照片筛选”，可选择快速模式或智能模式进入后续引导。", width_cm=15.8)
    add_heading(doc, "7.2 示例 B：CTF 估计与曝光筛选的决策支持", 2)
    add_body(doc, "用户可以问：“CTF fit 质量差时先检查什么？”系统会把问题绑定到当前检查点，给出顺序化排查：确认输入是否为运动校正后的 micrographs，检查 defocus 搜索范围与 amplitude contrast，查看 CTF fit 与冰层/污染情况，并提醒进行 Manual Curate。该回答是指导和检查清单，不会替用户直接修改 cryoSPARC/RELION 作业。")
    add_table(doc, ["输入", "系统响应", "用户仍需完成的判断"], [
        ("文字描述：CTF fit 差、曲线不稳定", "返回可能原因、检查顺序、回退路径和 QC 提醒", "结合原始图像、采集条件和实际 CTF 曲线确认"),
        ("上传 CTF 截图（Vision 已配置）", "尝试识别图中曲线/标签，并与检查点上下文合并回答", "核对识别结果；不能把图像识别当作最终测量"),
        ("未配置模型服务", "使用本地流程卡片、故障库和规则答案", "按软件界面和实验数据执行复核"),
    ], widths=[5.0, 6.5, 5.0])
    add_heading(doc, "7.3 示例 C：教学卡片与互动测验", 2)
    add_body(doc, "在教学模式下，用户可以阅读某一步的“做什么、关键参数含义、常见问题、判断标准”卡片，并完成题库中的选择题。测验的作用是帮助理解流程逻辑，不代表对实验结果的自动认证。")
    add_heading(doc, "7.4 示例 D：专家模式与工作流导出", 2)
    add_body(doc, "专家用户可以查看当前步骤的参数建议、规则依据和经验条目，保存个人预设，并导出 cryoSPARC Workflow JSON 作为后续配置起点。导出的 JSON 需要在目标 cryoSPARC 环境中检查版本、输入路径、资源和参数，StructPilot 不会自动提交作业。")
    add_source_note(doc, "图 4 为 2026 年 7 月 26 日本地运行实例截图；其余示例依据当前 UI 路由、知识库和导出功能整理。")
    doc.add_page_break()


def add_innovation(doc):
    add_heading(doc, "08 创新点、可靠性与隐私设计", 1)
    add_heading(doc, "8.1 创新点", 2)
    add_table(doc, ["创新点", "具体做法", "价值"], [
        ("检查点式陪跑", "以 12 个流程节点组织对话、QC 和状态", "把复杂流程变成可定位、可回溯的任务序列"),
        ("三模式同底座", "入门、教学、专家只改变信息呈现和交互密度", "同一知识体系覆盖学习、执行和复核"),
        ("本地优先的渐进式 AI", "规则/知识库是最低可用路径，模型为增强层", "降低部署门槛，也减少单一模型失效的影响"),
        ("经验治理", "支持实验室经验导入、审核、分级、废弃和追踪", "把个人经验转成团队可复用资产"),
        ("可导出而非黑箱自动化", "生成报告和 Workflow JSON，保留人工审批点", "兼顾效率、透明度与实验责任边界"),
    ], widths=[3.0, 8.4, 5.1])
    add_heading(doc, "8.2 可靠性与隐私", 2)
    add_bullets(doc, [
        "规则层与知识层优先：模型用于表达增强，不改变结构化结论的边界。",
        "输入和输出可追踪：会话、消息、检查点状态和部分元数据写入本地 SQLite。",
        "本地运行：演示数据和会话默认保存在本地；启用外部模型时，用户应根据机构要求审查发送内容。",
        "明确不确定性：参数推荐需要结合样品、采集条件、GPU 资源和软件版本进行专家复核。",
        "角色区分：当前支持 admin/member/guest 等角色概念，具体部署仍需结合本地访问控制策略。",
    ])
    add_callout(doc, "责任边界", "StructPilot 提供导航、解释、检查和配置建议；最终实验参数、数据质量判断、结构解析和论文级结论仍由研究人员负责。", fill="FFF4E5", accent=AMBER)
    doc.add_page_break()


def add_limits(doc):
    add_heading(doc, "09 部署要求、当前边界与后续路线", 1)
    add_heading(doc, "9.1 运行与部署要求", 2)
    add_table(doc, ["项目", "要求/说明"], [
        ("运行入口", "本地运行 Streamlit 应用；Windows 可使用项目提供的启动脚本"),
        ("Python", "Python 3.10+（以 requirements.txt 为准）"),
        ("数据存储", "SQLite 与本地 runtime 目录；适合演示和小团队使用"),
        ("模型服务", "LLM、Embedding、Vision、Speech 为可选配置，需提供相应 API/模型"),
        ("上传限制", "当前界面约束：图片 10 MB，音频 25 MB"),
        ("外部软件", "cryoSPARC、RELION 及 Coot/Phenix/ChimeraX 等由用户另行安装和运行"),
    ], widths=[4.2, 12.3])
    add_heading(doc, "9.2 当前边界（请评审按此理解）", 2)
    add_table(doc, ["类别", "当前边界"], [
        ("计算执行", "StructPilot 不执行 cryoSPARC/RELION 的重计算，也不自动提交作业。"),
        ("参数可信度", "参数是基于规则、知识和上下文的建议，不是对任何样品的保证值。"),
        ("模型依赖", "LLM、Vision、Embedding、Speech 的增强能力取决于外部配置、配额和网络。"),
        ("图像判断", "图像分析是辅助解释，必须由用户回看原图和软件 QC 结果。"),
        ("数据规模", "本地 SQLite 不是生产级多用户数据库，暂不宣称高并发和跨机构协作能力。"),
        ("结构建模", "模型构建与验证依赖外部专业工具，StructPilot 主要提供流程导航和验证提醒。"),
    ], widths=[4.0, 12.5])
    add_heading(doc, "9.3 后续路线", 2)
    add_numbered(doc, [
        "扩充 cp_02—cp_12 的教学卡片、故障案例和可验证 QC 规则。",
        "增加项目级数据隔离、审计日志和更适合团队部署的数据库后端。",
        "在明确授权和版本兼容策略后，探索与 cryoSPARC/RELION 的作业状态接口集成。",
        "建立基于真实实验案例的离线评测集，分别评估检索正确性、故障路径覆盖和建议可执行性。",
    ])
    add_callout(doc, "数据声明", "本文档不宣称视觉识别准确率、参数准确率、节省时间、用户满意度或结构求解性能等未经独立实验验证的指标。", fill=RED_LIGHT, accent=RED)
    doc.add_page_break()


def add_appendix(doc):
    add_heading(doc, "附录：功能核验表与建议演示脚本", 1)
    add_heading(doc, "A.1 功能核验表", 2)
    add_table(doc, ["核验项", "核验方式", "结论"], [
        ("版本与入口", "检查 version.py 与 Streamlit 启动入口", "StructPilot v6.0.0；Web UI 可启动"),
        ("流程导航", "选择软件、检查点和模式，输入开始/完成/跳过", "已实现"),
        ("会话持久化", "检查 SQLite sessions/messages/checkpoint 记录", "已实现"),
        ("知识库", "检查 flows、rules、faults、teaching、review 目录", "已实现"),
        ("多模态输入", "文本、图片上传/粘贴；语音路径需服务配置", "分层实现"),
        ("报告下载", "从界面导出报告/Markdown", "已实现"),
        ("cryoSPARC Workflow JSON", "专家模式导出并在目标环境复核", "已实现；不自动提交"),
        ("质量验证", "运行项目健康检查与核心测试", "已知核心测试通过；完整 pytest 仍有历史测试/临时目录问题，不作全部通过声明"),
    ], widths=[4.0, 8.5, 4.0])
    add_heading(doc, "A.2 5 分钟建议演示脚本", 2)
    add_numbered(doc, [
        "30 秒：展示定位——StructPilot 是流程导航与知识协作系统，不替代 cryoSPARC/RELION。",
        "60 秒：选择 cryoSPARC，展示步骤 3 的真实界面与快速/智能模式入口。",
        "90 秒：输入“CTF fit 质量差怎么办”，展示检查顺序、QC 提醒和当前检查点上下文。",
        "60 秒：切换教学模式，展示教学卡片与测验；说明同一底层状态不会改变。",
        "60 秒：切换专家模式，展示参数、经验、报告或 Workflow JSON 导出，并强调人工复核边界。",
        "30 秒：总结创新点——检查点、三模式、本地优先知识治理和可追踪输出。",
    ])
    add_heading(doc, "A.3 评审可复核的仓库位置", 2)
    add_table(doc, ["内容", "仓库位置"], [
        ("版本号", "version.py"),
        ("Web 主入口", "main.py / streamlit_app.py"),
        ("三种模式", "modes/beginner.py、modes/teaching.py、modes/expert.py"),
        ("流程定义", "knowledge_base/flows/pipeline_checkpoints.json"),
        ("规则与故障库", "knowledge_base/rules/、knowledge_base/faults/"),
        ("教学内容", "knowledge_base/teaching_cards.json、knowledge_base/quiz_bank.json"),
        ("审查治理", "knowledge_base/review/、knowledge_base/validate_kb_structure.py"),
        ("运行截图", "比赛提交文档/fig_live_home.png"),
    ], widths=[4.0, 12.5])
    add_source_note(doc, "本文档为决赛提交说明，不替代源代码、运行日志或实验记录；评审如需复核，以当前提交包中的代码和配置为准。")


def main():
    doc = Document()
    doc.core_properties.title = "StructPilot v6.0 决赛作品说明文档"
    doc.core_properties.subject = "冷冻电镜单颗粒分析智能导航系统的创意、功能与演示说明"
    doc.core_properties.author = "StructPilot 项目组"
    doc.core_properties.keywords = "StructPilot; Cryo-EM; cryoSPARC; RELION; 作品说明"
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_overview(doc)
    add_creative(doc)
    add_users(doc)
    add_functions(doc)
    add_modes(doc)
    add_architecture(doc)
    add_demo(doc)
    add_innovation(doc)
    add_limits(doc)
    add_appendix(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
